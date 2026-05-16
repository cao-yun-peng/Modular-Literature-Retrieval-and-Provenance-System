from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    input_path = root / "Manuscript.docx"
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    doc = Document(str(input_path))

    paragraphs: list[dict] = []
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").replace("\r", "").replace("\n", " ").strip()
        paragraphs.append({"i": i, "style": getattr(p.style, "name", None), "text": text})

    tables: list[dict] = []
    for ti, t in enumerate(doc.tables):
        rows: list[dict] = []
        for ri, row in enumerate(t.rows):
            cells: list[dict] = []
            for ci, cell in enumerate(row.cells):
                cell_text = (cell.text or "").replace("\r", "").replace("\n", " ").strip()
                cells.append({"c": ci, "text": cell_text})
            rows.append({"r": ri, "cells": cells})
        tables.append({"ti": ti, "rows": rows})

    extracted_txt_path = root / "Manuscript.extracted.txt"
    with extracted_txt_path.open("w", encoding="utf-8") as f:
        f.write("# Extracted from Manuscript.docx\n\n")
        f.write("## Paragraphs\n")
        for p in paragraphs:
            f.write(f"[P{p['i']:04d} | {p['style']}] {p['text']}\n")
        f.write("\n## Tables\n")
        for t in tables:
            f.write(f"[TABLE {t['ti']:03d}]\n")
            for row in t["rows"]:
                f.write("| " + " | ".join(cell["text"] for cell in row["cells"]) + " |\n")
            f.write("\n")

    stats = {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "nonempty_paragraphs": sum(1 for p in paragraphs if p["text"]),
        "nonempty_tables": sum(
            1
            for t in tables
            if any(any(c["text"] for c in r["cells"]) for r in t["rows"])
        ),
        "paragraph_characters": sum(len(p["text"]) for p in paragraphs),
        "table_characters": sum(len(c["text"]) for t in tables for r in t["rows"] for c in r["cells"]),
    }
    (root / "Manuscript.stats.json").write_text(
        json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print("Wrote", extracted_txt_path)
    print("Stats", json.dumps(stats, ensure_ascii=False))


if __name__ == "__main__":
    main()

from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    input_path = root / "Manuscript.docx"
    output_path = root / "翻译.docx"

    doc = Document(str(input_path))

    # Only replace paragraphs that have actual text in the extracted view.
    translations: dict[int, str] = {
        0: "Manuscript",
        3: (
            "In this study, we reveal how nonreciprocal interactions influence topological defects. "
            "Starting from an extended two-dimensional nematic model, we introduce a nonreciprocal coupling "
            "in which the interaction between a single director and its neighborhood depends asymmetrically on the "
            "state of that director itself: regardless of the orientation of neighboring directors, the interaction is "
            "enhanced along the director’s own orientation. Our results show that nonreciprocal interactions induce "
            "active-matter-like macroscopic behavior, thereby reshaping the annihilation dynamics and the coarsening process."
        ),
        4: (
            "The nonreciprocal two-dimensional nematic model consists of directors (orientational units) placed on the sites "
            "of a triangular lattice with linear size L. The local state of each director (unit) is described by an orientation "
            "angle $\\theta_i$ that obeys head–tail symmetry (i.e., $\\pi$ periodicity), and the corresponding physical order "
            "parameter is the $Q$ tensor. The system evolves according to:"
        ),
        6: (
            "Here, $\\gamma$ is the damping coefficient, $J$ is the coupling strength, $\\eta$ is Gaussian white noise with zero mean "
            "and unit variance, and $T$ is the bath temperature (with $k_B=1$ fixed). Nonreciprocity is introduced by weighting the "
            "coupling between neighboring spins (e.g., spins $i$ and $j$) according to the orientation of spin $i$ relative to the bond "
            "direction connecting the two spins (denoted as $\\varphi_{ij}$; see Fig. 1(a) for a schematic). This mechanism is encoded by a "
            "kernel function $g_\\sigma$. When $g_\\sigma$ is a constant, the model reduces to the reciprocal two-dimensional XY model with "
            "overdamped (nonconserved) dynamics [9]; a “vision-cone” interaction corresponds to a step-function kernel [31,48,53,54]. "
            "In this work, we adopt a smooth vision-cone kernel based on the von Mises distribution, $g_\\sigma(\\varphi)=\\exp(\\sigma\\cos 2\\varphi)$ [55], "
            "to mitigate discretization effects induced by the lattice geometry. The factor $\\cos(2\\varphi)$ reflects the head–tail symmetry, i.e., the director "
            "can “see” other directors both in the forward and backward directions. For small $\\varphi$, $g_\\sigma(\\varphi) \\sim e^{(-2\\varphi^2/(\\sigma^{-1}))}$; "
            "thus, $\\sigma$ effectively plays the role of the inverse variance (larger $\\sigma$ corresponds to a narrower vision cone)."
        ),
        19: (
            "In the two-dimensional nematic model, point-like topological defects carry half-integer charge $q$, given by the winding number [4]. "
            "The local director field around a $q=1/2$ defect located at the origin can be written as $\\theta(x,y)=q\\,\\mathrm{atan}(y/x)+\\mu$, where the integration constant "
            "$\\mu$ has a clear physical meaning as a shape parameter. For a $+1/2$ defect, $\\mu$ also indicates the direction of its motion (we should find an appropriate place to discuss "
            "the active motion of $+1/2$ defects, e.g., in the Introduction)."
        ),
        30: (
            "Annihilation. This work clarifies the central role of nonreciprocity in the dynamics of topological defects. In a hexagonal-lattice nematic at $T=0$, "
            "nonreciprocal interactions generate self-propulsion, endowing $+1/2$ defects with directed motility. For a pair consisting of a self-propelled $+1/2$ defect and a $-1/2$ defect, "
            "the terminal relaxed state depends sensitively on the initial condition. The physical origin of this sensitivity can be traced to a motion-induced “zipper” selection mechanism: the propulsion direction "
            "of the self-driven defect, together with the line connecting it to its partner defect, defines a degenerate orientation of the order parameter transverse to the trajectory. Like a zipper slider, the self-propelled defect continuously "
            "separates and expands, in the region ahead of its path, two candidate structures in the transverse direction behind the path (a nontrivial spin-wave configuration and a trivial zero-field configuration). Consequently, the initial orientation of the spin-wave configuration "
            "relative to the propulsion direction uniquely determines the macroscopic phase that the zipper mechanism selects and amplifies to the entire system."
        ),
        35: (
            "Under the coupling between nonreciprocal interactions and the hexagonal lattice geometry, we observe the splitting of a $+1$ topological defect. Theoretical analysis indicates that, in a hexagonal lattice, when one adopts a construction with minimal winding number 3, "
            "the maximum topological charge that a nematic can accommodate is $\\pm 0.5$. Therefore, the initial state of a $+1$ defect is, in essence, composed of one $-1/2$ defect and three $+1/2$ defects. By examining the splitting process, we find that in the absence of nonreciprocity, "
            "one of the $+1/2$ defects annihilates with the central $-1/2$ defect. However, under nonreciprocal interactions, the $+1/2$ and $-1/2$ defects do not exhibit the expected annihilation; instead, they display a distinct dynamical evolution: the central $-1/2$ defect remains stable, while the three $+1/2$ defects separate along the self-propulsion direction and gradually move away from the center. This behavior goes beyond the traditional theoretical framework of defect dynamics and reveals a nontrivial evolution of topological defects in nonreciprocal systems."
        ),
        49: (
            "To systematically investigate this phenomenon, we construct a phase diagram to determine the range of initial configurations for which the three $+1/2$ defects in the initial state can stably separate and move away along the self-propulsion direction under different nonreciprocity strengths. "
            "The results show that, as the nonreciprocity parameter $\\sigma$ increases, the set of initial configurations that exhibit this special dynamical behavior expands."
        ),
        54: (
            "In addition, we design different boundary conditions to further tune the defect dynamics, including a fixed boundary condition and an initial-shape boundary condition, and we examine the influence of different boundary shapes. We find that the boundary shape does not affect defect creation and annihilation; instead, it primarily regulates defect trajectories. "
            "This is because boundary-mediated interactions modulate the defect dynamics. During evolution, defects are jointly influenced by the self-propulsion induced by nonreciprocity, defect–defect interactions, and boundary effects. For a small system size under fixed boundary conditions, the $+1/2$ defects first move away from the central $-1/2$ defect and then gradually return due to boundary effects; eventually, one $+1/2$ defect annihilates with the central $-1/2$ defect, leaving only two $+1/2$ defects that undergo rotational motion. "
            "In contrast, under the initial-shape boundary condition, the three $+1/2$ defects can stably move along periodic orbits."
        ),
    }

    for idx, text in translations.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"Paragraph index {idx} out of range (len={len(doc.paragraphs)})")
        # Avoid overwriting non-text objects (e.g., equations) in empty paragraphs.
        # We only touch the indices we extracted as having text.
        doc.paragraphs[idx].text = text

    doc.save(str(output_path))
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
